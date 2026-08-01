import os
import fitz  # PyMuPDF - renders PDF pages as images
from google import genai
from google.genai import types
from app.core.config import GEMINI_API_KEY
from app.core.logging_config import get_logger

logger = get_logger(__name__)

# Gemini model names change over time - if this ever errors out,
# check the current list at ai.google.dev/gemini-api/docs/models
GEMINI_VISION_MODEL = "gemini-2.5-flash"


def _get_genai_client() -> genai.Client:
    """Helper to initialize genai.Client with API key check."""
    if not GEMINI_API_KEY or GEMINI_API_KEY == "your_gemini_api_key_here":
        raise ValueError(
            "GEMINI_API_KEY is missing or set to its placeholder value. "
            "A real key is required to OCR scanned PDF files."
        )
    return genai.Client(api_key=GEMINI_API_KEY)


def extract_text(file_path: str) -> str:
    """
    Reads a file and returns its text content as a single string.
    .pdf files are treated as scanned/image-based and go through Gemini's
    vision capability instead of a traditional OCR engine.
    """
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    elif extension == ".docx":
        return extract_text_from_docx(file_path)

    elif extension == ".pdf":
        # Smart check: try reading embedded text first before using Gemini Vision OCR
        with fitz.open(file_path) as doc:
            pages = [page.get_text("text").strip() for page in doc]
        direct_text = "\n".join(p for p in pages if p)
        if direct_text:
            return direct_text

        # Fallback to Gemini Vision OCR if no embedded text is present
        return extract_text_from_scanned_pdf(file_path)

    else:
        raise ValueError(f"Unsupported file type: {extension}")


def extract_text_from_docx(file_path: str) -> str:
    """Extracts text paragraphs and table cells from a .docx file."""
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx.") from exc

    doc = DocxDocument(file_path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_scanned_pdf(file_path: str) -> str:
    """
    Renders each PDF page as an image, then asks Gemini to read the
    text out of that image - the same way a person would read a
    photo of a document.
    """
    client = _get_genai_client()
    doc = fitz.open(file_path)
    full_text = ""

    for page in doc:
        pix = page.get_pixmap(dpi=200)
        image_bytes = pix.tobytes("png")

        response = client.models.generate_content(
            model=GEMINI_VISION_MODEL,
            contents=[
                "Extract all the text from this image exactly as it appears. "
                "Do not summarize, translate, or explain - output only the raw text.",
                types.Part.from_bytes(data=image_bytes, mime_type="image/png"),
            ],
        )
        full_text += (response.text or "") + "\n"

    return full_text
