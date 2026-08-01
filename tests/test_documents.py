"""
Tests for POST /api/v1/documents/upload

Covers:
- Unsupported file type → 400
- Empty file content → 400
- Valid .txt upload (mocked embed + FAISS) → 200
"""
import io
from unittest.mock import patch
import pytest
from docx import Document as DocxDocument

from app.services.document_processor import extract_text


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _txt_upload(client, content: bytes, filename: str = "test.txt"):
    return client.post(
        "/api/v1/documents/upload",
        files={"file": (filename, io.BytesIO(content), "text/plain")},
    )


def _pdf_upload(client, content: bytes, filename: str = "test.pdf"):
    return client.post(
        "/api/v1/documents/upload",
        files={"file": (filename, io.BytesIO(content), "application/pdf")},
    )


def _docx_upload(client, text: str, filename: str = "test.docx"):
    document = DocxDocument()
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                filename,
                buffer,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )


# ---------------------------------------------------------------------------
# Validation tests — no mocking needed, errors fire before any API call
# ---------------------------------------------------------------------------

def test_upload_rejects_unsupported_extension(client):
    r = _txt_upload(client, b"some content", filename="report.xlsx")
    assert r.status_code == 400
    assert ".xlsx" in r.json()["detail"]


def test_upload_rejects_png(client):
    r = client.post(
        "/api/v1/documents/upload",
        files={"file": ("photo.png", io.BytesIO(b"\x89PNG\r\n"), "image/png")},
    )
    assert r.status_code == 400


def test_upload_rejects_empty_txt(client):
    r = _txt_upload(client, b"   ")  # whitespace only
    assert r.status_code == 400
    assert "No readable text" in r.json()["detail"]


# ---------------------------------------------------------------------------
# Happy-path test — mock embed_chunks and add_chunks so no Gemini key needed
# ---------------------------------------------------------------------------

@patch("app.api.v1.documents.add_chunks")
@patch("app.api.v1.documents.embed_chunks", return_value=[[0.1] * 768])
def test_upload_txt_success(mock_embed, mock_add, client):
    """
    A valid .txt file goes through the full pipeline and returns
    a 200 with id, filename, chunks_created, and status fields.
    """
    content = b"ContextFlow AI is a RAG-powered document knowledge assistant."
    r = _txt_upload(client, content, filename="policy.txt")

    assert r.status_code == 200
    body = r.json()
    assert body["filename"] == "policy.txt"
    assert body["chunks_created"] >= 1
    assert body["status"] == "uploaded and indexed"
    assert "id" in body

    # Confirm both service functions were actually called
    mock_embed.assert_called_once()
    mock_add.assert_called_once()


@patch("app.api.v1.documents.add_chunks")
@patch("app.api.v1.documents.embed_chunks", return_value=[[0.1] * 768])
def test_upload_creates_db_record(mock_embed, mock_add, client):
    """After upload, the response includes a DB-generated id, proving the record was saved."""
    content = b"Company leave policy: 18 days paid leave per year."
    r = _txt_upload(client, content, filename="leave.txt")
    assert r.status_code == 200
    body = r.json()
    assert isinstance(body["id"], int)
    assert body["id"] >= 1
    assert body["filename"] == "leave.txt"


@patch("app.api.v1.documents.add_chunks")
@patch("app.api.v1.documents.embed_chunks", side_effect=Exception("Gemini down"))
def test_upload_returns_502_when_embedding_fails(mock_embed, mock_add, client):
    """If the embedding API fails, the endpoint returns 502 — not a crash."""
    content = b"Some document content here."
    r = _txt_upload(client, content)
    assert r.status_code == 502
    assert "process and index" in r.json()["detail"]


@patch("app.api.v1.documents.add_chunks")
@patch("app.api.v1.documents.embed_chunks", return_value=[[0.1] * 768])
def test_upload_docx_success(mock_embed, mock_add, client):
    response = _docx_upload(client, "The company supports remote work.")
    assert response.status_code == 200
    assert response.json()["filename"] == "test.docx"
    mock_embed.assert_called_once()
    mock_add.assert_called_once()


def test_extract_text_reads_docx_paragraphs_and_tables(tmp_path):
    document = DocxDocument()
    document.add_paragraph("Remote work is available two days per week.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Team"
    table.cell(0, 1).text = "Engineering"
    path = tmp_path / "policy.docx"
    document.save(path)

    extracted = extract_text(str(path))
    assert "Remote work" in extracted
    assert "Team | Engineering" in extracted


@patch("app.api.v1.documents.add_chunks")
@patch("app.api.v1.documents.embed_chunks", return_value=[[0.1] * 768])
def test_list_documents_returns_uploaded_documents(mock_embed, mock_add, client):
    upload_response = _txt_upload(client, b"A searchable policy.", filename="policy.txt")
    assert upload_response.status_code == 200

    response = client.get("/api/v1/documents")
    assert response.status_code == 200
    documents = response.json()
    assert len(documents) == 1
    assert documents[0]["id"] == upload_response.json()["id"]
    assert documents[0]["filename"] == "policy.txt"
    assert "uploaded_at" in documents[0]


@patch("app.api.v1.documents.remove_document_chunks", return_value=1)
@patch("app.api.v1.documents.add_chunks")
@patch("app.api.v1.documents.embed_chunks", return_value=[[0.1] * 768])
def test_delete_document_removes_record_and_index(
    mock_embed,
    mock_add,
    mock_remove,
    client,
):
    upload_response = _txt_upload(client, b"A document to remove.", filename="remove-me.txt")
    document_id = upload_response.json()["id"]

    response = client.delete(f"/api/v1/documents/{document_id}")
    assert response.status_code == 200
    assert response.json()["status"] == "deleted"
    assert response.json()["chunks_removed"] == 1
    assert response.json()["file_deleted"] is True
    mock_remove.assert_called_once()
    assert client.get("/api/v1/documents").json() == []


def test_delete_document_returns_404_for_unknown_id(client):
    response = client.delete("/api/v1/documents/99999")
    assert response.status_code == 404
